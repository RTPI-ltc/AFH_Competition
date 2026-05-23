from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Callable

from agent.rag.config import SUPPORTED_EXTENSIONS, TEXT_EXTENSIONS


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _load_pdf_bytes(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf 未安装；安装后可解析 PDF") from exc
    reader = PdfReader(io.BytesIO(data))
    pieces: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            pieces.append(text)
    return "\n\n".join(pieces)


def _load_docx_bytes(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装；安装后可解析 DOCX") from exc
    document = Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _load_text_bytes(data: bytes) -> str:
    return _decode_text_bytes(data)


_LOADER_BYTES: dict[str, Callable[[bytes], str]] = {
    ".pdf": _load_pdf_bytes,
    ".docx": _load_docx_bytes,
}


def supported_extension(ext: str) -> bool:
    return ext.lower() in SUPPORTED_EXTENSIONS


def load_file(path: Path) -> str:
    ext = path.suffix.lower()
    if not supported_extension(ext):
        raise ValueError(f"不支持的文件类型: {ext}")
    if ext in TEXT_EXTENSIONS:
        text = _read_text_file(path)
        if ext == ".json":
            return _pretty_json(text)
        return text
    loader = _LOADER_BYTES.get(ext)
    if loader is None:
        raise ValueError(f"无对应解析器: {ext}")
    return loader(path.read_bytes())


def load_bytes(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if not supported_extension(ext):
        raise ValueError(f"不支持的文件类型: {ext}")
    if ext in TEXT_EXTENSIONS:
        text = _load_text_bytes(data)
        if ext == ".json":
            return _pretty_json(text)
        return text
    loader = _LOADER_BYTES.get(ext)
    if loader is None:
        raise ValueError(f"无对应解析器: {ext}")
    return loader(data)


def _pretty_json(raw: str) -> str:
    try:
        payload = json.loads(raw)
    except Exception:
        return raw
    return json.dumps(payload, ensure_ascii=False, indent=2)


def iter_files(root: Path) -> list[Path]:
    if root.is_file():
        return [root] if supported_extension(root.suffix) else []
    if not root.exists():
        return []
    matched: list[Path] = []
    for path in root.rglob("*"):
        if path.is_file() and supported_extension(path.suffix):
            matched.append(path)
    return sorted(matched)
